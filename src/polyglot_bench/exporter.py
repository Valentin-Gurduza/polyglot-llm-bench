"""
Polyglot-LLM-Bench — Results Exporter

Generates evaluation CSVs, leaderboard tables, raw JSON results,
interactive HTML visual reports, and clean Markdown review documents.
Supports both per-model separate files and consolidated benchmark runs.
"""

from __future__ import annotations

import html
import json
import re
from pathlib import Path
from typing import Any

import pandas as pd

from polyglot_bench.models import (
    BenchmarkRunMeta,
    EvaluationRow,
    LeaderboardEntry,
    ModelResponse,
)


def _ensure_dir(path: Path) -> None:
    """Ensure the parent directory exists."""
    path.parent.mkdir(parents=True, exist_ok=True)


def sanitize_model_slug(model: str) -> str:
    """
    Convert model slug like 'openai/gpt-4o' or 'minimax/minimax-m3:free'
    to a clean, filesystem-safe filename component.
    """
    safe = model.replace("/", "__").replace(":", "_").replace(".", "_")
    safe = re.sub(r"[^\w\-]", "_", safe)
    return safe


# ---------------------------------------------------------------------------
# JSON Results (Consolidated & Per-Model)
# ---------------------------------------------------------------------------
def export_results_json(
    responses: list[ModelResponse],
    meta: BenchmarkRunMeta,
    path: str | Path,
) -> Path:
    """
    Export benchmark results as a consolidated JSON.

    Includes run metadata and all individual responses.
    """
    path = Path(path)
    _ensure_dir(path)

    output = {
        "meta": meta.model_dump(mode="json"),
        "responses": [r.model_dump(mode="json") for r in responses],
    }

    with open(path, "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, ensure_ascii=False, default=str)

    return path


def export_per_model_results_json(
    responses: list[ModelResponse],
    meta: BenchmarkRunMeta,
    output_dir: str | Path,
) -> list[Path]:
    """
    Export separate JSON results files for each evaluated model.

    File pattern: {output_dir}/results_{safe_model_name}.json
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Group responses by model
    model_groups: dict[str, list[ModelResponse]] = {}
    for r in responses:
        model_groups.setdefault(r.model, []).append(r)

    created_paths: list[Path] = []
    for model_name, model_responses in model_groups.items():
        safe_name = sanitize_model_slug(model_name)
        model_file = out_dir / f"results_{safe_name}.json"

        # Create model-specific metadata
        successful = sum(1 for r in model_responses if r.success)
        model_meta = meta.model_copy(
            update={
                "models": [model_name],
                "total_tasks": len(model_responses),
                "successful_tasks": successful,
                "failed_tasks": len(model_responses) - successful,
            }
        )

        output = {
            "meta": model_meta.model_dump(mode="json"),
            "model": model_name,
            "responses": [r.model_dump(mode="json") for r in model_responses],
        }

        with open(model_file, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, ensure_ascii=False, default=str)

        created_paths.append(model_file)

    return created_paths


# ---------------------------------------------------------------------------
# Evaluation Sheet CSV (Consolidated & Per-Model)
# ---------------------------------------------------------------------------
def _create_evaluation_dataframe(responses: list[ModelResponse]) -> pd.DataFrame:
    """Helper to convert model responses into an evaluation DataFrame."""
    rows: list[dict[str, Any]] = []
    for r in responses:
        row = EvaluationRow(
            prompt_id=r.prompt_id,
            language=r.language,
            category=r.category,
            model=r.model,
            prompt_text=r.prompt_text,
            model_response=r.raw_response if r.success else f"[ERROR] {r.error}",
            constraints="",
            constraint_adherence=None,
            linguistic_naturalness=None,
            factual_accuracy=None,
            tone_clarity=None,
            total_score=None,
            annotator_notes="",
        )
        rows.append(row.model_dump())

    df = pd.DataFrame(rows)
    column_order = [
        "prompt_id",
        "language",
        "category",
        "model",
        "prompt_text",
        "model_response",
        "constraints",
        "constraint_adherence",
        "linguistic_naturalness",
        "factual_accuracy",
        "tone_clarity",
        "total_score",
        "annotator_notes",
    ]
    return df[[c for c in column_order if c in df.columns]]


def export_evaluation_sheet(
    responses: list[ModelResponse],
    path: str | Path,
) -> Path:
    """
    Generate a consolidated evaluation_sheet.csv pre-populated with model outputs
    and empty scoring columns for human annotators.
    """
    path = Path(path)
    _ensure_dir(path)

    df = _create_evaluation_dataframe(responses)
    df.to_csv(path, index=False, encoding="utf-8")
    return path


def export_per_model_evaluation_sheets(
    responses: list[ModelResponse],
    output_dir: str | Path,
) -> list[Path]:
    """
    Export separate evaluation CSV sheets for each evaluated model.

    File pattern: {output_dir}/evaluation_sheet_{safe_model_name}.csv
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Group responses by model
    model_groups: dict[str, list[ModelResponse]] = {}
    for r in responses:
        model_groups.setdefault(r.model, []).append(r)

    created_paths: list[Path] = []
    for model_name, model_responses in model_groups.items():
        safe_name = sanitize_model_slug(model_name)
        model_file = out_dir / f"evaluation_sheet_{safe_name}.csv"

        df = _create_evaluation_dataframe(model_responses)
        df.to_csv(model_file, index=False, encoding="utf-8")
        created_paths.append(model_file)

    return created_paths


# ---------------------------------------------------------------------------
# Clean Markdown to HTML Formatter (Python Fallback & Pre-renderer)
# ---------------------------------------------------------------------------
def _format_markdown_to_clean_html(raw_text: str) -> str:
    """
    Converts raw markdown with asterisks (**bold**), headers (###), lists, etc.
    into clean, readable HTML paragraphs and elements.
    """
    if not raw_text:
        return ""

    escaped = html.escape(raw_text)

    # 1. Code blocks: ```lang ... ```
    def _code_block_sub(match: re.Match) -> str:
        code_content = match.group(2)
        return f'<div class="md-code-block"><pre><code>{code_content}</code></pre></div>'

    escaped = re.sub(r"```([a-zA-Z0-9_\-\.]*)\n?(.*?)```", _code_block_sub, escaped, flags=re.DOTALL)

    # 2. Inline code: `code`
    escaped = re.sub(r"`([^`]+)`", r'<code class="inline-code">\1</code>', escaped)

    # 3. Headers: ###, ##, #
    escaped = re.sub(r"^###\s+(.*?)$", r'<h4 class="md-h4">\1</h4>', escaped, flags=re.MULTILINE)
    escaped = re.sub(r"^##\s+(.*?)$", r'<h3 class="md-h3">\1</h3>', escaped, flags=re.MULTILINE)
    escaped = re.sub(r"^#\s+(.*?)$", r'<h2 class="md-h2">\1</h2>', escaped, flags=re.MULTILINE)

    # 4. Bold: **text** or __text__
    escaped = re.sub(r"\*\*(.+?)\*\*", r'<strong>\1</strong>', escaped)
    escaped = re.sub(r"__(.+?)__", r'<strong>\1</strong>', escaped)

    # 5. Italic: *text* or _text_ (single asterisk not inside strong)
    escaped = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r'<em>\1</em>', escaped)
    escaped = re.sub(r"(?<!_)_([^_]+?)_(?!_)", r'<em>\1</em>', escaped)

    # 6. Blockquotes: > text
    escaped = re.sub(r"^>\s+(.*?)$", r'<blockquote class="md-quote">\1</blockquote>', escaped, flags=re.MULTILINE)

    # 7. Convert lines to structured paragraphs and list items
    lines = escaped.split("\n")
    processed_blocks: list[str] = []
    current_para: list[str] = []
    in_list = False
    list_type = "ul"

    def _flush_para() -> None:
        nonlocal current_para
        if current_para:
            p_content = "<br>".join(current_para).strip()
            if p_content:
                processed_blocks.append(f'<p class="md-p">{p_content}</p>')
            current_para = []

    def _flush_list() -> None:
        nonlocal in_list, list_type
        if in_list:
            processed_blocks.append(f"</{list_type}>")
            in_list = False

    for line in lines:
        stripped = line.strip()

        # Check for start/end of code blocks or headers or quotes
        if line.startswith("<div class=\"md-code-block\">") or line.startswith("<h") or line.startswith("<blockquote"):
            _flush_para()
            _flush_list()
            processed_blocks.append(line)
            continue

        if not stripped:
            _flush_para()
            _flush_list()
            continue

        # Bullet list item: - or *
        bullet_match = re.match(r"^[-*•]\s+(.*)$", stripped)
        if bullet_match:
            _flush_para()
            if not in_list or list_type != "ul":
                _flush_list()
                processed_blocks.append('<ul class="md-list">')
                in_list = True
                list_type = "ul"
            processed_blocks.append(f"<li>{bullet_match.group(1)}</li>")
            continue

        # Numbered list item: 1. or 1)
        num_match = re.match(r"^(\d+)[\.\)]\s+(.*)$", stripped)
        if num_match:
            _flush_para()
            if not in_list or list_type != "ol":
                _flush_list()
                processed_blocks.append('<ol class="md-list">')
                in_list = True
                list_type = "ol"
            processed_blocks.append(f"<li>{num_match.group(2)}</li>")
            continue

        # Regular text line
        _flush_list()
        current_para.append(stripped)

    _flush_para()
    _flush_list()

    return "\n".join(processed_blocks)


# ---------------------------------------------------------------------------
# Visual HTML Evaluation Report (Per-Model & Master)
# ---------------------------------------------------------------------------
def _generate_model_html(model_name: str, responses: list[ModelResponse]) -> str:
    """Generate a responsive, beautifully styled HTML review dossier for a model with interactive 1-5 scoring rubric."""
    total = len(responses)
    successful = sum(1 for r in responses if r.success)
    failed = total - successful
    avg_latency = (
        sum(r.latency_ms for r in responses if r.success) / successful
        if successful
        else 0
    )
    total_tokens = sum(r.token_usage.total_tokens for r in responses if r.success)

    cards_html = []
    tasks_data = []

    for i, r in enumerate(responses, 1):
        status_badge = (
            '<span class="badge badge-success">✓ Success</span>'
            if r.success
            else '<span class="badge badge-error">✗ Failed</span>'
        )
        lang_badge = f'<span class="badge badge-lang">{html.escape(r.language.upper())}</span>'
        cat_badge = f'<span class="badge badge-cat">{html.escape(r.category.replace("_", " ").title())}</span>'
        latency_str = f"{r.latency_ms:.0f} ms" if r.latency_ms else "N/A"
        tokens_str = f"{r.token_usage.completion_tokens} tokens" if r.token_usage.completion_tokens else ""

        if r.success:
            formatted_html = _format_markdown_to_clean_html(r.raw_response)
            raw_escaped = html.escape(r.raw_response)
        else:
            formatted_html = f'<div class="error-msg">⚠️ Request Failed: {html.escape(r.error or "Unknown error")}</div>'
            raw_escaped = html.escape(r.error or "Unknown error")

        tasks_data.append({
            "prompt_id": r.prompt_id,
            "language": r.language,
            "category": r.category,
            "model": r.model,
            "prompt_text": r.prompt_text,
            "model_response": r.raw_response if r.success else f"[ERROR] {r.error}",
        })

        card = f"""
        <div class="eval-card" id="task-{r.prompt_id}">
            <div class="card-header">
                <div class="card-title">
                    <span class="task-num">Task #{r.prompt_id}</span>
                    {lang_badge}
                    {cat_badge}
                </div>
                <div class="card-meta">
                    {status_badge}
                    <span class="meta-item">⏱ {latency_str}</span>
                    {f'<span class="meta-item">🔤 {tokens_str}</span>' if tokens_str else ''}
                    <button type="button" class="toggle-btn" onclick="toggleRawView({r.prompt_id})">
                        <span id="btn-text-{r.prompt_id}">📄 Raw Output</span>
                    </button>
                </div>
            </div>

            <div class="card-body">
                <div class="section-label">📌 Prompt Instruction & Constraints:</div>
                <div class="prompt-box">{html.escape(r.prompt_text)}</div>

                <div class="section-label">🤖 Model Response:</div>
                
                <!-- Clean Formatted Output (Default) -->
                <div class="rendered-response" id="rendered-{r.prompt_id}">
                    {formatted_html}
                </div>

                <!-- Raw Monospace View (Hidden by default) -->
                <div class="raw-response" id="raw-{r.prompt_id}" style="display: none;">
                    <pre><code>{raw_escaped}</code></pre>
                </div>

                <!-- Human Evaluation Rubric (1-5 Scale) -->
                <div class="scoring-section">
                    <div class="scoring-header">
                        <span class="section-label" style="margin-bottom: 0;">✍️ Human Evaluator Rubric (1–5 Scale)</span>
                        <span class="score-summary-pill" id="score-summary-{r.prompt_id}">Score: <strong>-- / 20</strong></span>
                    </div>

                    <div class="rubric-grid">
                        <div class="rubric-card">
                            <div class="rubric-title">
                                <span>1. Constraint Adherence</span>
                                <span class="rubric-val" id="val-ca-{r.prompt_id}">--</span>
                            </div>
                            <div class="rubric-desc">Sentence count, word limit, forbidden words, schema</div>
                            <div class="score-btns" data-task="{r.prompt_id}" data-dim="ca">
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ca', 1)">1</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ca', 2)">2</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ca', 3)">3</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ca', 4)">4</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ca', 5)">5</button>
                            </div>
                        </div>

                        <div class="rubric-card">
                            <div class="rubric-title">
                                <span>2. Linguistic Naturalness</span>
                                <span class="rubric-val" id="val-ln-{r.prompt_id}">--</span>
                            </div>
                            <div class="rubric-desc">Grammar, diacritics (ă/î/ș/ț), fluency, idiomatic style</div>
                            <div class="score-btns" data-task="{r.prompt_id}" data-dim="ln">
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ln', 1)">1</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ln', 2)">2</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ln', 3)">3</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ln', 4)">4</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'ln', 5)">5</button>
                            </div>
                        </div>

                        <div class="rubric-card">
                            <div class="rubric-title">
                                <span>3. Factual Accuracy</span>
                                <span class="rubric-val" id="val-fa-{r.prompt_id}">--</span>
                            </div>
                            <div class="rubric-desc">Dates, historical facts, calculation accuracy</div>
                            <div class="score-btns" data-task="{r.prompt_id}" data-dim="fa">
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'fa', 1)">1</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'fa', 2)">2</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'fa', 3)">3</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'fa', 4)">4</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'fa', 5)">5</button>
                            </div>
                        </div>

                        <div class="rubric-card">
                            <div class="rubric-title">
                                <span>4. Tone & Clarity</span>
                                <span class="rubric-val" id="val-tc-{r.prompt_id}">--</span>
                            </div>
                            <div class="rubric-desc">Appropriate formality, audience localization, helpfulness</div>
                            <div class="score-btns" data-task="{r.prompt_id}" data-dim="tc">
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'tc', 1)">1</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'tc', 2)">2</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'tc', 3)">3</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'tc', 4)">4</button>
                                <button type="button" class="btn-score" onclick="setScore({r.prompt_id}, 'tc', 5)">5</button>
                            </div>
                        </div>
                    </div>

                    <div class="notes-row">
                        <input type="text" id="note-{r.prompt_id}" class="notes-field" placeholder="Annotator critique & notes for Task #{r.prompt_id}..." oninput="saveNote({r.prompt_id})">
                    </div>
                </div>
            </div>
        </div>
        """
        cards_html.append(card)

    cards_str = "\n".join(cards_html)
    tasks_json_str = json.dumps(tasks_data, ensure_ascii=False)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Benchmark Evaluation — {html.escape(model_name)}</title>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-color: #0b0f19;
            --card-bg: #151d30;
            --card-header-bg: #1a233a;
            --border-color: #263352;
            --text-main: #f8fafc;
            --text-body: #e2e8f0;
            --text-muted: #94a3b8;
            --accent-blue: #38bdf8;
            --accent-cyan: #22d3ee;
            --accent-purple: #c084fc;
            --accent-green: #4ade80;
            --accent-red: #f87171;
            --prompt-bg: #0d1322;
            --response-bg: #101728;
            --code-bg: #070a12;
            --border-highlight: #3b82f6;
            --rubric-bg: #0d1424;
            --rubric-card-bg: #131b2e;
        }}
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            line-height: 1.7;
            padding: 2.5rem 1rem;
        }}
        .container {{
            max-width: 1060px;
            margin: 0 auto;
        }}
        header {{
            background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 2.25rem;
            margin-bottom: 2.5rem;
            box-shadow: 0 10px 30px rgba(0, 0, 0, 0.4);
        }}
        .header-title {{
            font-size: 1.85rem;
            font-weight: 700;
            color: var(--text-main);
            margin-bottom: 0.4rem;
            letter-spacing: -0.02em;
        }}
        .header-subtitle {{
            color: var(--accent-cyan);
            font-family: 'JetBrains Mono', monospace;
            font-size: 1.05rem;
            font-weight: 500;
            margin-bottom: 1.5rem;
        }}
        .header-actions {{
            display: flex;
            gap: 1rem;
            margin-bottom: 1.75rem;
            flex-wrap: wrap;
        }}
        .action-btn {{
            background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%);
            color: #ffffff;
            border: none;
            border-radius: 8px;
            padding: 0.65rem 1.25rem;
            font-size: 0.9rem;
            font-weight: 600;
            cursor: pointer;
            display: inline-flex;
            align-items: center;
            gap: 0.5rem;
            transition: all 0.2s ease;
            box-shadow: 0 4px 12px rgba(37, 99, 235, 0.3);
        }}
        .action-btn:hover {{
            background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
            transform: translateY(-1px);
        }}
        .action-btn.secondary {{
            background: rgba(30, 41, 59, 0.8);
            border: 1px solid var(--border-color);
            box-shadow: none;
            color: var(--text-body);
        }}
        .action-btn.secondary:hover {{
            background: rgba(51, 65, 85, 0.8);
            color: #ffffff;
        }}
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(170px, 1fr));
            gap: 1.25rem;
        }}
        .stat-box {{
            background: rgba(15, 23, 42, 0.7);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 1.15rem;
            text-align: center;
        }}
        .stat-val {{
            font-size: 1.6rem;
            font-weight: 700;
            color: var(--text-main);
        }}
        .stat-lbl {{
            font-size: 0.8rem;
            color: var(--text-muted);
            text-transform: uppercase;
            letter-spacing: 0.06em;
            margin-top: 0.35rem;
            font-weight: 600;
        }}
        .eval-card {{
            background-color: var(--card-bg);
            border: 1px solid var(--border-color);
            border-radius: 14px;
            margin-bottom: 2.25rem;
            overflow: hidden;
            box-shadow: 0 6px 20px rgba(0, 0, 0, 0.25);
            transition: transform 0.15s ease, border-color 0.2s ease;
        }}
        .eval-card:hover {{
            border-color: var(--border-highlight);
        }}
        .card-header {{
            background: var(--card-header-bg);
            border-bottom: 1px solid var(--border-color);
            padding: 1.1rem 1.4rem;
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            gap: 0.85rem;
        }}
        .card-title {{
            display: flex;
            align-items: center;
            gap: 0.75rem;
        }}
        .task-num {{
            font-weight: 700;
            font-size: 1.15rem;
            color: var(--accent-cyan);
            font-family: 'JetBrains Mono', monospace;
        }}
        .card-meta {{
            display: flex;
            align-items: center;
            gap: 0.85rem;
            font-size: 0.85rem;
        }}
        .meta-item {{
            color: var(--text-muted);
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.85rem;
        }}
        .badge {{
            display: inline-block;
            padding: 0.28rem 0.65rem;
            border-radius: 6px;
            font-size: 0.75rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }}
        .badge-lang {{
            background: #1e3a8a;
            color: #bfdbfe;
            border: 1px solid #3b82f6;
        }}
        .badge-cat {{
            background: #4c1d95;
            color: #e9d5ff;
            border: 1px solid #8b5cf6;
        }}
        .badge-success {{
            background: rgba(34, 197, 94, 0.15);
            color: var(--accent-green);
            border: 1px solid var(--accent-green);
        }}
        .badge-error {{
            background: rgba(239, 68, 68, 0.15);
            color: var(--accent-red);
            border: 1px solid var(--accent-red);
        }}
        .toggle-btn {{
            background: rgba(56, 189, 248, 0.1);
            color: var(--accent-blue);
            border: 1px solid rgba(56, 189, 248, 0.3);
            border-radius: 6px;
            padding: 0.35rem 0.75rem;
            font-size: 0.8rem;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.2s ease;
        }}
        .toggle-btn:hover {{
            background: rgba(56, 189, 248, 0.2);
            border-color: var(--accent-blue);
            color: #ffffff;
        }}
        .card-body {{
            padding: 1.5rem;
        }}
        .section-label {{
            font-size: 0.85rem;
            font-weight: 700;
            color: var(--text-muted);
            margin-bottom: 0.6rem;
            text-transform: uppercase;
            letter-spacing: 0.06em;
        }}
        .prompt-box {{
            background-color: var(--prompt-bg);
            border: 1px solid var(--border-color);
            border-radius: 10px;
            padding: 1.15rem 1.3rem;
            font-size: 0.95rem;
            color: #cbd5e1;
            margin-bottom: 1.5rem;
            white-space: pre-wrap;
            line-height: 1.65;
        }}
        
        /* Clean Rendered Response View */
        .rendered-response {{
            background-color: var(--response-bg);
            border: 1px solid var(--border-color);
            border-radius: 10px;
            padding: 1.5rem 1.75rem;
            color: var(--text-body);
            font-size: 1rem;
            line-height: 1.75;
            margin-bottom: 1.75rem;
        }}
        .rendered-response .md-p {{
            margin-bottom: 1rem;
        }}
        .rendered-response .md-p:last-child {{
            margin-bottom: 0;
        }}
        .rendered-response strong {{
            color: #ffffff;
            font-weight: 700;
        }}
        .rendered-response em {{
            color: var(--accent-cyan);
            font-style: italic;
        }}
        .rendered-response .md-h2 {{
            font-size: 1.35rem;
            font-weight: 700;
            color: var(--accent-cyan);
            margin: 1.25rem 0 0.6rem 0;
        }}
        .rendered-response .md-h3 {{
            font-size: 1.18rem;
            font-weight: 700;
            color: var(--accent-blue);
            margin: 1.1rem 0 0.5rem 0;
        }}
        .rendered-response .md-h4 {{
            font-size: 1.05rem;
            font-weight: 600;
            color: var(--text-main);
            margin: 1rem 0 0.4rem 0;
        }}
        .rendered-response .md-list {{
            margin: 0.75rem 0 1rem 1.75rem;
        }}
        .rendered-response .md-list li {{
            margin-bottom: 0.45rem;
            padding-left: 0.35rem;
        }}
        .rendered-response .md-quote {{
            border-left: 4px solid var(--accent-blue);
            padding: 0.6rem 1rem;
            margin: 1rem 0;
            background: rgba(56, 189, 248, 0.05);
            border-radius: 0 8px 8px 0;
            color: #cbd5e1;
            font-style: italic;
        }}
        .rendered-response .inline-code {{
            background: var(--code-bg);
            color: #f472b6;
            font-family: 'JetBrains Mono', monospace;
            padding: 0.15rem 0.45rem;
            border-radius: 5px;
            font-size: 0.88em;
            border: 1px solid rgba(244, 114, 182, 0.2);
        }}
        .rendered-response .md-code-block {{
            background: var(--code-bg);
            border: 1px solid var(--border-color);
            border-radius: 8px;
            padding: 1.15rem;
            margin: 1rem 0;
            overflow-x: auto;
        }}
        .rendered-response .md-code-block pre {{
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.9rem;
            color: #f8fafc;
            line-height: 1.5;
        }}
        
        /* Raw Monospace View */
        .raw-response {{
            background-color: var(--code-bg);
            border: 1px solid var(--border-color);
            border-radius: 10px;
            padding: 1.25rem;
            overflow-x: auto;
            margin-bottom: 1.75rem;
        }}
        .raw-response pre {{
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.88rem;
            color: #f1f5f9;
            white-space: pre-wrap;
            word-break: break-word;
        }}

        /* Interactive Scoring Section */
        .scoring-section {{
            background: var(--rubric-bg);
            border: 1px solid var(--border-color);
            border-radius: 12px;
            padding: 1.5rem;
            margin-top: 1rem;
        }}
        .scoring-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 1.25rem;
            flex-wrap: wrap;
            gap: 0.5rem;
        }}
        .score-summary-pill {{
            background: #1e293b;
            border: 1px solid var(--border-color);
            padding: 0.35rem 0.9rem;
            border-radius: 20px;
            font-size: 0.9rem;
            font-weight: 600;
            color: var(--text-muted);
        }}
        .score-summary-pill strong {{
            color: var(--accent-cyan);
            font-size: 1rem;
        }}
        .rubric-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
            gap: 1rem;
            margin-bottom: 1.25rem;
        }}
        .rubric-card {{
            background: var(--rubric-card-bg);
            border: 1px solid rgba(255, 255, 255, 0.06);
            border-radius: 10px;
            padding: 1rem;
        }}
        .rubric-title {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            font-weight: 600;
            font-size: 0.88rem;
            color: var(--text-main);
            margin-bottom: 0.25rem;
        }}
        .rubric-val {{
            font-family: 'JetBrains Mono', monospace;
            font-weight: 700;
            color: var(--accent-cyan);
            font-size: 0.95rem;
        }}
        .rubric-desc {{
            font-size: 0.75rem;
            color: var(--text-muted);
            line-height: 1.4;
            margin-bottom: 0.75rem;
            min-height: 32px;
        }}
        .score-btns {{
            display: flex;
            gap: 0.4rem;
        }}
        .btn-score {{
            flex: 1;
            background: rgba(255, 255, 255, 0.05);
            color: var(--text-body);
            border: 1px solid rgba(255, 255, 255, 0.1);
            border-radius: 6px;
            padding: 0.4rem 0;
            font-size: 0.85rem;
            font-weight: 700;
            cursor: pointer;
            transition: all 0.15s ease;
        }}
        .btn-score:hover {{
            background: rgba(56, 189, 248, 0.2);
            color: var(--accent-blue);
            border-color: var(--accent-blue);
        }}
        .btn-score.active {{
            background: #2563eb;
            color: #ffffff;
            border-color: #3b82f6;
            box-shadow: 0 0 10px rgba(37, 99, 235, 0.5);
        }}
        .notes-row {{
            margin-top: 1rem;
        }}
        .notes-field {{
            width: 100%;
            background: #090d16;
            border: 1px solid var(--border-color);
            border-radius: 8px;
            padding: 0.75rem 1rem;
            color: var(--text-body);
            font-family: 'Inter', sans-serif;
            font-size: 0.88rem;
            outline: none;
            transition: border-color 0.2s ease;
        }}
        .notes-field:focus {{
            border-color: var(--accent-blue);
        }}
        .error-msg {{
            color: var(--accent-red);
            font-weight: 600;
            padding: 0.75rem;
            background: rgba(239, 68, 68, 0.1);
            border-radius: 8px;
            border: 1px solid rgba(239, 68, 68, 0.3);
        }}
    </style>
</head>
<body>
    <div class="container">
        <header>
            <div class="header-title">LLM Benchmark Evaluation Report</div>
            <div class="header-subtitle">Model: {html.escape(model_name)}</div>
            
            <div class="header-actions">
                <button type="button" class="action-btn" onclick="exportAnnotatedCSV()">
                    💾 Export Scored CSV Sheet
                </button>
                <button type="button" class="action-btn secondary" onclick="resetAllScores()">
                    🔄 Reset Scoring Form
                </button>
            </div>

            <div class="stats-grid">
                <div class="stat-box">
                    <div class="stat-val">{total}</div>
                    <div class="stat-lbl">Total Prompts</div>
                </div>
                <div class="stat-box">
                    <div class="stat-val" style="color: var(--accent-green);">{successful}</div>
                    <div class="stat-lbl">Successful</div>
                </div>
                <div class="stat-box">
                    <div class="stat-val" style="color: {'var(--accent-red)' if failed else 'var(--text-muted)'};">{failed}</div>
                    <div class="stat-lbl">Failed</div>
                </div>
                <div class="stat-box">
                    <div class="stat-val">{avg_latency:.0f} ms</div>
                    <div class="stat-lbl">Avg Latency</div>
                </div>
                <div class="stat-box">
                    <div class="stat-val">{total_tokens:,}</div>
                    <div class="stat-lbl">Total Tokens</div>
                </div>
            </div>
        </header>

        <main>
            {cards_str}
        </main>
    </div>

    <script>
        const MODEL_SLUG = "{html.escape(sanitize_model_slug(model_name))}";
        const TASKS_DATA = {tasks_json_str};

        // Store active ratings in memory and sync with localStorage
        const scoresState = {{}};

        function loadSavedScores() {{
            const saved = localStorage.getItem(`polyglot_scores_${{MODEL_SLUG}}`);
            if (saved) {{
                try {{
                    const parsed = JSON.parse(saved);
                    Object.assign(scoresState, parsed);
                    for (const taskId in scoresState) {{
                        const tState = scoresState[taskId];
                        if (tState.ca) applyButtonUI(taskId, 'ca', tState.ca);
                        if (tState.ln) applyButtonUI(taskId, 'ln', tState.ln);
                        if (tState.fa) applyButtonUI(taskId, 'fa', tState.fa);
                        if (tState.tc) applyButtonUI(taskId, 'tc', tState.tc);
                        if (tState.notes) {{
                            const noteInput = document.getElementById(`note-${{taskId}}`);
                            if (noteInput) noteInput.value = tState.notes;
                        }}
                        updateTotalScoreUI(taskId);
                    }}
                }} catch (e) {{}}
            }}
        }}

        function setScore(taskId, dim, value) {{
            if (!scoresState[taskId]) scoresState[taskId] = {{}};
            scoresState[taskId][dim] = value;

            applyButtonUI(taskId, dim, value);
            updateTotalScoreUI(taskId);
            saveToStorage();
        }}

        function applyButtonUI(taskId, dim, value) {{
            const container = document.querySelector(`.score-btns[data-task="${{taskId}}"][data-dim="${{dim}}"]`);
            if (container) {{
                const btns = container.querySelectorAll('.btn-score');
                btns.forEach((btn, idx) => {{
                    if (idx + 1 === value) {{
                        btn.classList.add('active');
                    }} else {{
                        btn.classList.remove('active');
                    }}
                }});
            }}
            const valLabel = document.getElementById(`val-${{dim}}-${{taskId}}`);
            if (valLabel) valLabel.innerText = `${{value}} / 5`;
        }}

        function updateTotalScoreUI(taskId) {{
            const t = scoresState[taskId] || {{}};
            const ca = t.ca || 0;
            const ln = t.ln || 0;
            const fa = t.fa || 0;
            const tc = t.tc || 0;

            const totalSpan = document.getElementById(`score-summary-${{taskId}}`);
            if (!totalSpan) return;

            const scoredCount = (t.ca ? 1 : 0) + (t.ln ? 1 : 0) + (t.fa ? 1 : 0) + (t.tc ? 1 : 0);
            if (scoredCount === 0) {{
                totalSpan.innerHTML = `Score: <strong>-- / 20</strong>`;
            }} else {{
                const sum = ca + ln + fa + tc;
                totalSpan.innerHTML = `Score: <strong>${{sum}} / 20</strong> (${{scoredCount}}/4 rated)`;
            }}
        }}

        function saveNote(taskId) {{
            const noteInput = document.getElementById(`note-${{taskId}}`);
            if (!scoresState[taskId]) scoresState[taskId] = {{}};
            scoresState[taskId].notes = noteInput ? noteInput.value : "";
            saveToStorage();
        }}

        function saveToStorage() {{
            localStorage.setItem(`polyglot_scores_${{MODEL_SLUG}}`, JSON.stringify(scoresState));
        }}

        function resetAllScores() {{
            if (confirm("Reset all evaluator scores and notes for this model?")) {{
                localStorage.removeItem(`polyglot_scores_${{MODEL_SLUG}}`);
                location.reload();
            }}
        }}

        function toggleRawView(taskId) {{
            const rendered = document.getElementById(`rendered-${{taskId}}`);
            const raw = document.getElementById(`raw-${{taskId}}`);
            const btnText = document.getElementById(`btn-text-${{taskId}}`);

            if (rendered.style.display === "none") {{
                rendered.style.display = "block";
                raw.style.display = "none";
                btnText.innerText = "📄 Raw Output";
            }} else {{
                rendered.style.display = "none";
                raw.style.display = "block";
                btnText.innerText = "✨ Formatted Text";
            }}
        }}

        function exportAnnotatedCSV() {{
            const headers = [
                "prompt_id",
                "language",
                "category",
                "model",
                "prompt_text",
                "model_response",
                "constraint_adherence",
                "linguistic_naturalness",
                "factual_accuracy",
                "tone_clarity",
                "total_score",
                "annotator_notes"
            ];

            const rows = [headers];

            TASKS_DATA.forEach(task => {{
                const s = scoresState[task.prompt_id] || {{}};
                const ca = s.ca || "";
                const ln = s.ln || "";
                const fa = s.fa || "";
                const tc = s.tc || "";
                const total = (s.ca && s.ln && s.fa && s.tc) ? (s.ca + s.ln + s.fa + s.tc) : "";
                const notes = s.notes || "";

                rows.push([
                    task.prompt_id,
                    task.language,
                    task.category,
                    task.model,
                    task.prompt_text,
                    task.model_response,
                    ca,
                    ln,
                    fa,
                    tc,
                    total,
                    notes
                ]);
            }});

            const csvContent = rows.map(r => r.map(cell => `"${{String(cell).replace(/"/g, '""')}}"`).join(",")).join("\\n");
            const blob = new Blob([csvContent], {{ type: 'text/csv;charset=utf-8;' }});
            const link = document.createElement("a");
            const url = URL.createObjectURL(blob);
            link.setAttribute("href", url);
            link.setAttribute("download", `evaluation_sheet_${{MODEL_SLUG}}_annotated.csv`);
            document.body.appendChild(link);
            link.click();
            document.body.removeChild(link);
        }}

        document.addEventListener("DOMContentLoaded", () => {{
            loadSavedScores();
        }});
    </script>
</body>
</html>
"""


def export_per_model_html_reports(
    responses: list[ModelResponse],
    output_dir: str | Path,
) -> list[Path]:
    """
    Export visual, beautifully formatted HTML review files for each evaluated model.

    File pattern: {output_dir}/report_{safe_model_name}.html
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Group responses by model
    model_groups: dict[str, list[ModelResponse]] = {}
    for r in responses:
        model_groups.setdefault(r.model, []).append(r)

    created_paths: list[Path] = []
    for model_name, model_responses in model_groups.items():
        safe_name = sanitize_model_slug(model_name)
        report_file = out_dir / f"report_{safe_name}.html"

        html_content = _generate_model_html(model_name, model_responses)
        with open(report_file, "w", encoding="utf-8") as f:
            f.write(html_content)

        created_paths.append(report_file)

    return created_paths


# ---------------------------------------------------------------------------
# Markdown Evaluation Dossier (Per-Model)
# ---------------------------------------------------------------------------
def _generate_model_markdown(model_name: str, responses: list[ModelResponse]) -> str:
    """Generate a clean markdown report easily viewable in VS Code, Obsidian, Notion."""
    successful = sum(1 for r in responses if r.success)
    lines = [
        f"# 📋 Evaluation Dossier: `{model_name}`\n",
        f"- **Total Prompts Evaluated**: {len(responses)}",
        f"- **Successful Completions**: {successful}/{len(responses)}",
        "- **Evaluation Framework**: Polyglot-LLM-Bench\n",
        "---\n",
    ]

    for r in responses:
        status_str = "✅ Success" if r.success else "❌ Failed"
        lines.append(f"## Prompt #{r.prompt_id} [{r.language.upper()} / {r.category}] — {status_str}\n")
        lines.append(f"**Latency**: {r.latency_ms:.0f} ms | **Tokens**: {r.token_usage.completion_tokens}\n")
        lines.append("### 📌 Prompt Instruction")
        lines.append(f"> {r.prompt_text}\n")
        lines.append("### 🤖 Model Response")
        if r.success:
            lines.append("```text")
            lines.append(r.raw_response)
            lines.append("```\n")
        else:
            lines.append(f"**Error**: `{r.error}`\n")

        lines.append("### ✍️ Evaluation Scores (1–5)")
        lines.append("| Constraint Adherence | Linguistic Naturalness | Factual Accuracy | Tone & Clarity | Total (4–20) |")
        lines.append("|:---:|:---:|:---:|:---:|:---:|")
        lines.append("| [ ] / 5 | [ ] / 5 | [ ] / 5 | [ ] / 5 | [ ] / 20 |\n")
        lines.append("**Annotator Notes**:\n")
        lines.append("---\n")

    return "\n".join(lines)


def export_per_model_markdown_reports(
    responses: list[ModelResponse],
    output_dir: str | Path,
) -> list[Path]:
    """
    Export Markdown dossiers for each evaluated model.

    File pattern: {output_dir}/report_{safe_model_name}.md
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Group responses by model
    model_groups: dict[str, list[ModelResponse]] = {}
    for r in responses:
        model_groups.setdefault(r.model, []).append(r)

    created_paths: list[Path] = []
    for model_name, model_responses in model_groups.items():
        safe_name = sanitize_model_slug(model_name)
        report_file = out_dir / f"report_{safe_name}.md"

        md_content = _generate_model_markdown(model_name, model_responses)
        with open(report_file, "w", encoding="utf-8") as f:
            f.write(md_content)

        created_paths.append(report_file)

    return created_paths


# ---------------------------------------------------------------------------
# Word Document Report (.docx) (Per-Model)
# ---------------------------------------------------------------------------
def _generate_model_docx(model_name: str, responses: list[ModelResponse], doc_path: Path) -> Path:
    """Generate a clean, styled Microsoft Word (.docx) review document for a model."""
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = docx.Document()

    # Set standard margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Polyglot-LLM-Bench Evaluation Report")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 58, 138)  # Deep Navy Blue

    subtitle_p = doc.add_paragraph()
    sub_run = subtitle_p.add_run(f"Model: {model_name}\n")
    sub_run.font.size = Pt(12)
    sub_run.font.bold = True
    sub_run.font.color.rgb = RGBColor(71, 85, 105)

    # Summary Statistics Table
    total = len(responses)
    successful = sum(1 for r in responses if r.success)
    failed = total - successful
    avg_latency = (
        sum(r.latency_ms for r in responses if r.success) / successful
        if successful
        else 0
    )
    total_tokens = sum(r.token_usage.total_tokens for r in responses if r.success)

    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_headers = ["Total Tasks", "Successful", "Failed", "Avg Latency", "Total Tokens"]
    summary_values = [str(total), str(successful), str(failed), f"{avg_latency:.0f} ms", f"{total_tokens:,}"]

    for col_idx, (hdr, val) in enumerate(zip(summary_headers, summary_values)):
        hdr_cell = summary_table.cell(0, col_idx)
        hdr_cell.text = hdr
        hdr_cell.paragraphs[0].runs[0].font.bold = True
        hdr_cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        hdr_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Light blue background for headers
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E2E8F0"/>')
        hdr_cell._tc.get_or_add_tcPr().append(shading)

        val_cell = summary_table.cell(1, col_idx)
        val_cell.text = val
        val_cell.paragraphs[0].runs[0].font.bold = True
        val_cell.paragraphs[0].runs[0].font.size = Pt(11)
        val_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # Spacing

    # Section divider
    doc.add_heading("Detailed Task Outputs & Human Evaluation", level=1)

    for r in responses:
        # Heading for each task
        status_text = "✓ Success" if r.success else "✗ Failed"
        task_heading = doc.add_heading(
            f"Task #{r.prompt_id} [{r.language.upper()} / {r.category.replace('_', ' ').title()}] — {status_text}",
            level=2,
        )

        # Meta paragraph
        meta_p = doc.add_paragraph()
        meta_run = meta_p.add_run(
            f"Latency: {r.latency_ms:.0f} ms  |  Prompt Tokens: {r.token_usage.prompt_tokens}  |  Completion Tokens: {r.token_usage.completion_tokens}"
        )
        meta_run.font.size = Pt(8.5)
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor(100, 116, 139)

        # Prompt Box
        doc.add_heading("📌 Prompt Instruction", level=3)
        prompt_table = doc.add_table(rows=1, cols=1)
        prompt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        prompt_cell = prompt_table.cell(0, 0)
        prompt_cell.text = r.prompt_text
        prompt_cell.paragraphs[0].runs[0].font.size = Pt(10)
        prompt_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
        prompt_cell._tc.get_or_add_tcPr().append(prompt_shading)

        doc.add_paragraph()

        # Model Response Box
        doc.add_heading("🤖 Model Output", level=3)
        resp_table = doc.add_table(rows=1, cols=1)
        resp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        resp_cell = resp_table.cell(0, 0)
        resp_cell.text = r.raw_response if r.success else f"[ERROR]: {r.error}"
        for p in resp_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = "Consolas"
        resp_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
        resp_cell._tc.get_or_add_tcPr().append(resp_shading)

        doc.add_paragraph()

        # Human Scoring Rubric Table
        doc.add_heading("✍️ Evaluator Scoring Rubric (1–5)", level=3)
        rubric_table = doc.add_table(rows=2, cols=6)
        rubric_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rubric_headers = [
            "Constraint (1-5)",
            "Fluency (1-5)",
            "Accuracy (1-5)",
            "Tone (1-5)",
            "Total (4-20)",
            "Notes / Critique",
        ]
        for c_idx, rh in enumerate(rubric_headers):
            rc = rubric_table.cell(0, c_idx)
            rc.text = rh
            rc.paragraphs[0].runs[0].font.bold = True
            rc.paragraphs[0].runs[0].font.size = Pt(8.5)
            r_shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E2E8F0"/>')
            rc._tc.get_or_add_tcPr().append(r_shading)

            # Empty row for annotator
            val_rc = rubric_table.cell(1, c_idx)
            val_rc.text = "[   ]" if c_idx < 5 else ""
            if len(val_rc.paragraphs[0].runs) > 0:
                val_rc.paragraphs[0].runs[0].font.size = Pt(9)
                val_rc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing between tasks

    doc.save(str(doc_path))
    return doc_path


def export_per_model_docx_reports(
    responses: list[ModelResponse],
    output_dir: str | Path,
) -> list[Path]:
    """
    Export styled Microsoft Word documents (.docx) for each evaluated model.

    File pattern: {output_dir}/report_{safe_model_name}.docx
    """
    out_dir = Path(output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    # Group responses by model
    model_groups: dict[str, list[ModelResponse]] = {}
    for r in responses:
        model_groups.setdefault(r.model, []).append(r)

    created_paths: list[Path] = []
    for model_name, model_responses in model_groups.items():
        safe_name = sanitize_model_slug(model_name)
        doc_file = out_dir / f"report_{safe_name}.docx"

        _generate_model_docx(model_name, model_responses, doc_file)
        created_paths.append(doc_file)

    return created_paths


# ---------------------------------------------------------------------------
# Leaderboard CSV
# ---------------------------------------------------------------------------
def export_leaderboard(
    responses: list[ModelResponse],
    path: str | Path,
    *,
    accumulate_existing: bool = True,
) -> Path:
    """
    Generate a leaderboard.csv with aggregated per-model-per-language stats.

    If accumulate_existing is True and path exists, updates existing models
    or appends new model rows.
    """
    path = Path(path)
    _ensure_dir(path)

    # Group by (model, language)
    groups: dict[tuple[str, str], list[ModelResponse]] = {}
    for r in responses:
        key = (r.model, r.language)
        groups.setdefault(key, []).append(r)

    new_entries: list[dict[str, Any]] = []
    for (model, language), resps in sorted(groups.items()):
        successful = [r for r in resps if r.success]
        entry = LeaderboardEntry(
            model=model,
            language=language,
            total_prompts=len(resps),
            successful=len(successful),
            failed=len(resps) - len(successful),
            avg_latency_ms=(
                sum(r.latency_ms for r in successful) / len(successful)
                if successful
                else 0.0
            ),
            avg_prompt_tokens=(
                sum(r.token_usage.prompt_tokens for r in successful) / len(successful)
                if successful
                else 0.0
            ),
            avg_completion_tokens=(
                sum(r.token_usage.completion_tokens for r in successful) / len(successful)
                if successful
                else 0.0
            ),
        )
        new_entries.append(entry.model_dump())

    new_df = pd.DataFrame(new_entries)

    if accumulate_existing and path.exists():
        try:
            existing_df = pd.read_csv(path)
            # Combine and drop duplicates on (model, language), keeping the latest run
            combined_df = pd.concat([existing_df, new_df], ignore_index=True)
            combined_df = combined_df.drop_duplicates(subset=["model", "language"], keep="last")
            combined_df.to_csv(path, index=False, encoding="utf-8")
            return path
        except Exception:
            pass

    new_df.to_csv(path, index=False, encoding="utf-8")
    return path

